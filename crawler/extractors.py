"""Turn raw bytes of various file types into title + text + links."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

from bs4 import BeautifulSoup

from .utils import normalize_url

# Optional heavy extractors -- degrade gracefully if missing OR broken. We catch
# BaseException because a misconfigured native dependency (e.g. cryptography
# pulled in by pypdf) can raise non-Exception errors at import time, and an
# optional extractor must never be able to take down the crawler.
try:
    from pypdf import PdfReader
except BaseException:  # pragma: no cover
    PdfReader = None

try:
    import docx  # python-docx
except BaseException:  # pragma: no cover
    docx = None


_WHITESPACE = re.compile(r"\s+")


@dataclass
class Document:
    url: str
    title: str = ""
    text: str = ""
    content_type: str = ""
    links: list[str] = field(default_factory=list)


def _clean(text: str) -> str:
    return _WHITESPACE.sub(" ", text).strip()


def extract(url: str, content_type: str, body: bytes) -> Document:
    """Dispatch on content type and return a populated Document."""
    ct = (content_type or "").lower()

    if ct in ("text/html", "application/xhtml+xml") or _looks_like_html(body, ct):
        return _extract_html(url, body, ct or "text/html")
    if ct == "application/pdf":
        return _extract_pdf(url, body)
    if ct == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return _extract_docx(url, body)
    if ct.startswith("text/"):
        return Document(
            url=url,
            content_type=ct,
            text=_clean(body.decode("utf-8", errors="replace")),
        )

    # Unknown / binary: no text, no links -- caller may still store metadata.
    return Document(url=url, content_type=ct or "application/octet-stream")


def _looks_like_html(body: bytes, ct: str) -> bool:
    if ct and not ct.startswith("text/"):
        return False
    head = body[:512].lstrip().lower()
    return head.startswith(b"<!doctype html") or head.startswith(b"<html")


def _extract_html(url: str, body: bytes, ct: str) -> Document:
    soup = BeautifulSoup(body, "lxml")

    for tag in soup(["script", "style", "noscript", "template"]):
        tag.decompose()

    title = ""
    if soup.title and soup.title.string:
        title = _clean(soup.title.string)
    if not title and soup.h1:
        title = _clean(soup.h1.get_text())

    text = _clean(soup.get_text(separator=" "))

    links: list[str] = []
    seen: set[str] = set()
    for a in soup.find_all("a", href=True):
        normalized = normalize_url(a["href"], base=url)
        if normalized and normalized not in seen:
            seen.add(normalized)
            links.append(normalized)

    return Document(url=url, title=title, text=text, content_type=ct, links=links)


def _extract_pdf(url: str, body: bytes) -> Document:
    if PdfReader is None:
        return Document(url=url, content_type="application/pdf")
    try:
        reader = PdfReader(io.BytesIO(body))
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        meta_title = ""
        if reader.metadata and reader.metadata.title:
            meta_title = _clean(str(reader.metadata.title))
        return Document(
            url=url,
            title=meta_title,
            text=_clean(" ".join(parts)),
            content_type="application/pdf",
        )
    except Exception:
        return Document(url=url, content_type="application/pdf")


def _extract_docx(url: str, body: bytes) -> Document:
    ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if docx is None:
        return Document(url=url, content_type=ct)
    try:
        document = docx.Document(io.BytesIO(body))
        text = _clean(" ".join(p.text for p in document.paragraphs))
        title = _clean(document.paragraphs[0].text) if document.paragraphs else ""
        return Document(url=url, title=title, text=text, content_type=ct)
    except Exception:
        return Document(url=url, content_type=ct)
